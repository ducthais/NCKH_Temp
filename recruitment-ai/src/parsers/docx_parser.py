import docx
from pathlib import Path
from typing import Dict, Any

def parse_docx_native(file_path: str | Path) -> Dict[str, Any]:
    """
    Trích xuất text từ file .docx
    """
    try:
        doc = docx.Document(str(file_path))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
                
        # Handle tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
                        
        raw_text = "\n".join(full_text)
        
        return {
            "status": "success",
            "raw_text": raw_text,
            "method": "docx2txt",
            "num_pages": 1 # DOCX doesn't easily expose page count without rendering
        }
    except Exception as e:
        return {
            "status": "error",
            "message": str(e),
            "raw_text": ""
        }
